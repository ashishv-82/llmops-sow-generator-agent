"""
Document handler utilities for PDF/DOCX parsing and export.
"""

from pathlib import Path

from docx import Document
from pypdf import PdfReader


def parse_document(file_path: str | Path) -> str:
    """
    Extract text from a PDF or DOCX file.

    Args:
        file_path: Path to the document file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file format is not supported
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return _parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Supported: .pdf, .docx")


def _parse_pdf(file_path: Path) -> str:
    """Extract text from PDF file."""
    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        text_parts.append(page.extract_text())

    return "\n\n".join(text_parts)


def _parse_docx(file_path: Path) -> str:
    """Extract text from DOCX file."""
    doc = Document(str(file_path))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    return "\n\n".join(text_parts)


def export_to_docx(content: str, output_path: str | Path) -> None:
    """
    Export markdown/text content to a DOCX file.

    Args:
        content: Text content (markdown will be rendered simply)
        output_path: Path for the output DOCX file
    """
    output_path = Path(output_path)

    # Create document
    doc = Document()

    # Split by markdown headers and paragraphs
    lines = content.split("\n")
    current_para: list[str] = []

    for line in lines:
        line = line.strip()

        if not line:
            # Empty line - add paragraph if we have content
            if current_para:
                doc.add_paragraph(" ".join(current_para))
                current_para = []
            continue

        # Check for headers
        if line.startswith("# "):
            if current_para:
                doc.add_paragraph(" ".join(current_para))
                current_para = []
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            if current_para:
                doc.add_paragraph(" ".join(current_para))
                current_para = []
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            if current_para:
                doc.add_paragraph(" ".join(current_para))
                current_para = []
            doc.add_heading(line[4:], level=3)
        else:
            # Regular text - clean markdown formatting
            cleaned = line.replace("**", "").replace("*", "")
            current_para.append(cleaned)

    # Add any remaining paragraph
    if current_para:
        doc.add_paragraph(" ".join(current_para))

    # Save document
    doc.save(str(output_path))
